"""
Summary Report Generator for Customer-Facing Documents
Generates concise, actionable summaries from detailed compliance reports
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional
import logging

logger = logging.getLogger(__name__)

class SummaryReportGenerator:
    """Generates concise summary reports for customers"""
    
    def __init__(self):
        self.brand_color = RGBColor(0, 123, 255)  # Professional blue
        self.accent_color = RGBColor(40, 167, 69)  # Success green
        self.warning_color = RGBColor(255, 193, 7)  # Warning amber
        
    def generate_summary(
        self,
        doc_data: Dict,
        evaluation: Dict,
        output_path: Optional[Path] = None
    ) -> str:
        """
        Generate a concise customer-facing summary report
        
        Args:
            doc_data: Extracted document data
            evaluation: Compliance evaluation results
            output_path: Where to save the summary
            
        Returns:
            Path to generated summary document
        """
        try:
            # Create new document
            doc = Document()
            
            # Set document margins
            for section in doc.sections:
                section.top_margin = Inches(0.8)
                section.bottom_margin = Inches(0.8)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Add header
            self._add_summary_header(doc, doc_data, evaluation)
            
            # Add key findings section
            self._add_key_findings(doc, evaluation)
            
            # Add priority actions
            self._add_priority_actions(doc, evaluation)
            
            # Add score visualization
            self._add_score_summary(doc, evaluation)
            
            # Add next steps
            self._add_next_steps(doc, doc_data, evaluation)
            
            # Add Overview and Actions sections (per client example format)
            self._add_overview_and_actions(doc, evaluation)
            
            # Add contact footer
            self._add_contact_footer(doc)
            
            # Save document
            if output_path is None:
                output_path = Path("results") / f"summary_{doc_data.get('filename', 'report')}"
            
            doc.save(str(output_path))
            logger.info(f"Summary report generated: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Error generating summary report: {e}")
            raise
    
    def _add_summary_header(self, doc: Document, doc_data: Dict, evaluation: Dict):
        """Add professional header to summary"""
        
        # Title
        title = doc.add_heading('Lake Management Analysis Summary', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Get document type from evaluation (hybrid analysis) or fallback to doc_data
        doc_type = evaluation.get('document_type', doc_data.get('document_type', 'hybrid'))
        type_breakdown = evaluation.get('type_breakdown', {})
        hybrid_analysis = evaluation.get('hybrid_analysis', {})
        characteristics = hybrid_analysis.get('document_characteristics', {})
        
        # Determine the display label
        plan_pct = type_breakdown.get('plan_percentage', characteristics.get('plan_percentage', 50))
        report_pct = type_breakdown.get('report_percentage', characteristics.get('report_percentage', 50))
        
        if report_pct > 70:
            doc_type_label = "Lake Report Analysis"
            doc_type_detail = f"(Primarily Retrospective Data - {report_pct}%)"
        elif plan_pct > 70:
            doc_type_label = "Lake Management Plan Analysis"
            doc_type_detail = f"(Primarily Forward-Looking - {plan_pct}%)"
        else:
            doc_type_label = "Hybrid Document Analysis"
            doc_type_detail = f"(Report: {report_pct}% | Plan: {plan_pct}%)"
        
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run(f"{doc_type_label}\n").bold = True
        info_para.add_run(f"{doc_type_detail}\n").italic = True
        info_para.add_run(f"{doc_data.get('filename', 'Document')}\n")
        info_para.add_run(f"{datetime.now().strftime('%B %d, %Y')}")
        
        doc.add_paragraph()  # Spacing
    
    def _add_key_findings(self, doc: Document, evaluation: Dict):
        """Add key findings section"""
        
        doc.add_heading('Key Findings', 1)
        
        # Overall score with interpretation
        score = evaluation.get('compliance_percentage', 0)
        # Handle compliance_level in various formats
        compliance_level = evaluation.get('compliance_level')
        if isinstance(compliance_level, str):
            # It's already a string value
            level = compliance_level
        elif hasattr(compliance_level, 'value'):
            # It's an Enum object
            level = compliance_level.value
        elif isinstance(compliance_level, dict):
            # It's a dictionary
            level = compliance_level.get('value', 'UNKNOWN')
        else:
            level = str(compliance_level) if compliance_level else 'UNKNOWN'
        
        # Score box
        score_para = doc.add_paragraph()
        score_para.add_run(f'Overall Compliance Score: ').bold = True
        score_run = score_para.add_run(f'{score:.0f}%')
        score_run.font.size = Pt(14)
        
        if score >= 70:
            score_run.font.color.rgb = self.accent_color
            interpretation = "Your report shows good understanding of lake management principles."
        elif score >= 50:
            score_run.font.color.rgb = self.warning_color
            interpretation = "Your report has a solid foundation but needs improvement in key areas."
        else:
            score_run.font.color.rgb = RGBColor(220, 53, 69)
            interpretation = "Your report is missing critical elements for effective lake management."
        
        doc.add_paragraph(interpretation)
        
        # Strengths section (clean format per client feedback - no checkmarks)
        all_strengths = evaluation.get('strengths', [])
        # Filter out partial credit items, duplicates, and clean up the text
        display_strengths = []
        seen_strengths = set()
        for s in all_strengths:
            if "(PARTIAL:" not in s:
                # Normalize for duplicate detection
                s_lower = s.lower().strip()
                if s_lower not in seen_strengths:
                    seen_strengths.add(s_lower)
                    display_strengths.append(s)
        
        if display_strengths:
            doc.add_heading('Strengths:', 2)
            for strength in display_strengths[:5]:
                doc.add_paragraph(f'{strength}', style='List Bullet')
        
        # Areas for Improvement section (clean format per client feedback)
        # Use a set to track unique items and avoid duplicates
        areas_for_improvement = []
        seen_items = set()
        
        # Add missing calculations (but avoid duplicates with similar names)
        # Also skip absolute calculations if percentage version is found
        found_calcs = evaluation.get('calculations', {}).get('found', [])
        found_calc_names = set()
        for fc in found_calcs:
            if isinstance(fc, dict):
                found_calc_names.add(fc.get('name', '').lower())
            else:
                found_calc_names.add(str(fc).lower())
        
        # Check if percentage versions are found (to avoid marking absolute as missing)
        has_pct_volume = any('percentage_volume' in n or 'percentage volume' in n for n in found_calc_names)
        has_pct_area = any('percentage_area' in n or 'percentage area' in n for n in found_calc_names)
        
        missing_calcs = evaluation.get('calculations', {}).get('missing', [])
        for calc in missing_calcs:
            if isinstance(calc, dict):
                calc_name = calc.get('name', '').replace('_', ' ')
            else:
                calc_name = str(calc).replace('_', ' ')
            
            # Skip if we already have a similar item
            calc_key = calc_name.lower()
            if calc_key not in seen_items:
                # Don't duplicate if percentage version is found
                if 'percentage' not in calc_key:
                    # Skip absolute volume if percentage volume is found
                    if 'water volume' in calc_key and has_pct_volume:
                        continue
                    # Skip absolute area if percentage area is found
                    if 'sediment area' in calc_key and has_pct_area:
                        continue
                    seen_items.add(calc_key)
                    areas_for_improvement.append(f"Missing calculation: {calc_name}")
        
        # Add weaknesses (avoiding duplicates)
        weaknesses = evaluation.get('weaknesses', [])
        for weakness in weaknesses:
            weakness_lower = weakness.lower().strip()
            # Skip if already covered or is a duplicate
            if weakness_lower not in seen_items:
                if "Missing critical parameter" not in weakness:  # Avoid duplicates with params
                    seen_items.add(weakness_lower)
                    areas_for_improvement.append(weakness)
        
        # Remove any remaining duplicates by converting to unique list
        unique_areas = []
        seen_text = set()
        for item in areas_for_improvement:
            item_normalized = item.lower().strip()
            if item_normalized not in seen_text:
                seen_text.add(item_normalized)
                unique_areas.append(item)
        
        if unique_areas:
            doc.add_heading('Areas for Improvement:', 2)
            for item in unique_areas[:5]:
                doc.add_paragraph(f'{item}', style='List Bullet')
        
        # Phytoplankton Measurement Assessment (per client feedback)
        self._add_phytoplankton_assessment(doc, evaluation)
        
        doc.add_paragraph()  # Spacing
    
    def _add_phytoplankton_assessment(self, doc: Document, evaluation: Dict):
        """Add phytoplankton measurement assessment section per client feedback"""
        
        # Always show this section to prompt the questions
        doc.add_paragraph("We need to add phytoplankton measurement")
        
        # Assessment questions (per client feedback)
        questions = [
            "Are you actually identifying phytoplankton types and species?",
            "Are you calculating the percentage of:",
        ]
        
        for question in questions:
            doc.add_paragraph(f'{question}', style='List Bullet')
        
        # Sub-items for percentage calculation
        sub_items = [
            ("Cyanobacteria", [
                "Toxin producers",
                "Non-Toxin producers"
            ]),
            ("Beneficial algae", [])
        ]
        
        for item, sub_list in sub_items:
            doc.add_paragraph(f'    ○ {item}')
            for sub in sub_list:
                doc.add_paragraph(f'        ■ {sub}')
        
        doc.add_paragraph('Are you measuring total phytoplankton biovolume?', style='List Bullet')
        
        # If phytoplankton management interventions are found, add concerns
        phyto_mgmt = evaluation.get('phytoplankton_management', {})
        if phyto_mgmt.get('is_negative') or phyto_mgmt.get('interventions_found'):
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run('Phytoplankton Management Concerns:').bold = True
            interventions = phyto_mgmt.get('interventions_found', [])
            if interventions:
                for intervention in interventions:
                    name = intervention.get('name', 'Unknown')
                    issue = intervention.get('issue', '')
                    doc.add_paragraph(f'{name}: {issue}', style='List Bullet')
                doc.add_paragraph(
                    "Note: Chemical interventions like algaecides, herbicides, and phosphorus "
                    "precipitants treat symptoms, not root causes."
                )
        
        # NEW: Add comprehensive analysis section if available
        self._add_comprehensive_analysis(doc, evaluation)
    
    def _add_lake_background(self, doc: Document, evaluation: Dict):
        """Add History & Background section per client feedback"""
        
        comp_analysis = evaluation.get('comprehensive_analysis', {})
        lake_bg = comp_analysis.get('lake_background', {})
        
        if not lake_bg or lake_bg.get('error'):
            return
        
        # Only show section if we have meaningful content
        has_content = (lake_bg.get('lake_name') or lake_bg.get('management_history') or 
                      lake_bg.get('current_systems') or lake_bg.get('stated_objectives'))
        
        if not has_content:
            return
        
        doc.add_heading('History & Background', 2)
        
        # Lake name and management history
        if lake_bg.get('lake_name'):
            p = doc.add_paragraph()
            p.add_run(f"Lake: ").bold = True
            p.add_run(lake_bg['lake_name'])
        
        if lake_bg.get('management_start_year'):
            doc.add_paragraph(f"Lake has been under management since {lake_bg['management_start_year']}.")
        
        if lake_bg.get('management_history'):
            doc.add_paragraph(lake_bg['management_history'])
        
        # Nutrient sources
        if lake_bg.get('nutrient_sources'):
            p = doc.add_paragraph()
            p.add_run('Identified Nutrient Sources: ').bold = True
            p.add_run(', '.join(lake_bg['nutrient_sources']))
        
        # Current systems
        if lake_bg.get('current_systems'):
            p = doc.add_paragraph()
            p.add_run('Current Treatment Systems: ').bold = True
            p.add_run(', '.join(lake_bg['current_systems']))
        
        # System issues
        if lake_bg.get('system_issues'):
            doc.add_paragraph('System Issues/Limitations:', style='List Bullet')
            for issue in lake_bg['system_issues'][:3]:
                doc.add_paragraph(f'  • {issue}')
        
        # Stated objectives
        if lake_bg.get('stated_objectives'):
            p = doc.add_paragraph()
            p.add_run('Stated Objectives: ').bold = True
            for obj in lake_bg['stated_objectives'][:3]:
                doc.add_paragraph(f'  • {obj}')
        
        # HAB history
        if lake_bg.get('hab_history'):
            p = doc.add_paragraph()
            p.add_run('HAB History: ').bold = True
            p.add_run(lake_bg['hab_history'])
        
        doc.add_paragraph()  # Spacing
    
    def _add_specific_values(self, doc: Document, evaluation: Dict):
        """Add section highlighting specific extracted values per client feedback"""
        
        comp_analysis = evaluation.get('comprehensive_analysis', {})
        specific = comp_analysis.get('specific_values', {})
        
        if not specific or specific.get('error'):
            return
        
        # Check if we have any meaningful specific values
        has_values = (specific.get('key_parameter_changes') or specific.get('do_values_by_year') or
                     specific.get('srp_values_by_year') or specific.get('hypoxic_percentage'))
        
        if not has_values:
            return
        
        doc.add_heading('Key Data Values', 2)
        
        # Deepest location info
        if specific.get('deepest_location_name'):
            p = doc.add_paragraph()
            p.add_run('Primary Sampling Location: ').bold = True
            p.add_run(specific['deepest_location_name'])
            doc.add_paragraph("(Deepest part of lake - most meaningful data)")
        
        # Show values by year if available
        if specific.get('do_values_by_year'):
            p = doc.add_paragraph()
            p.add_run('DO Values: ').bold = True
            p.add_run(specific['do_values_by_year'])
        
        if specific.get('srp_values_by_year'):
            p = doc.add_paragraph()
            p.add_run('SRP Values: ').bold = True
            p.add_run(specific['srp_values_by_year'])
        
        if specific.get('ammonia_values_by_year'):
            p = doc.add_paragraph()
            p.add_run('Ammonia Values: ').bold = True
            p.add_run(specific['ammonia_values_by_year'])
        
        # Year-over-year changes (most important per client)
        if specific.get('key_parameter_changes'):
            doc.add_heading('Year-over-Year Changes', 3)
            for change in specific['key_parameter_changes'][:5]:
                doc.add_paragraph(f'• {change}', style='List Bullet')
        
        # Hypoxic calculations if present
        if specific.get('hypoxic_percentage'):
            p = doc.add_paragraph()
            p.add_run('Hypoxic Water: ').bold = True
            p.add_run(f"{specific['hypoxic_percentage']:.1f}% of lake volume")
            if specific.get('hypoxic_volume_m3'):
                p.add_run(f" ({specific['hypoxic_volume_m3']:,.0f} m³)")
        
        if specific.get('biomass_potential_tonnes'):
            p = doc.add_paragraph()
            p.add_run('Phytoplankton Biomass Potential: ').bold = True
            p.add_run(f"{specific['biomass_potential_tonnes']:.1f} metric tonnes")
            doc.add_paragraph("(Potential algal bloom capacity based on available phosphorus)")
        
        doc.add_paragraph()  # Spacing
    
    def _add_data_quality_checklist(self, doc: Document, evaluation: Dict):
        """Add comprehensive data quality checklist per client's requirements"""
        
        comp_analysis = evaluation.get('comprehensive_analysis', {})
        checklist = comp_analysis.get('data_quality_checklist', {})
        
        if not checklist or checklist.get('error'):
            return
        
        doc.add_heading('Data Quality Assessment', 2)
        
        # Dissolved Oxygen Section
        doc.add_heading('Dissolved Oxygen', 3)
        
        # Create a simple checklist format
        do_items = [
            ('do_measured_at_intervals_to_bottom', 'Measured at regular intervals to bottom?'),
            ('do_multi_year_available', 'Multi-year data for trend analysis?'),
            ('oxycline_depth_identified', 'Oxycline depth (DO < 2.5 mg/L) identified?'),
            ('hypsographic_table_available', 'Hypsographic table available?'),
            ('hypoxic_volume_calculated', 'Hypoxic water volume calculated?'),
            ('hypoxic_area_calculated', 'Hypoxic sediment area calculated?'),
        ]
        
        for key, label in do_items:
            value = checklist.get(key, False)
            symbol = '✓' if value else '✗'
            color_word = '' if value else ' (MISSING)'
            doc.add_paragraph(f'{symbol} {label}{color_word}', style='List Bullet')
        
        # Show measurement frequency
        freq = checklist.get('do_measurement_frequency', 'unknown')
        doc.add_paragraph(f'   Measurement frequency: {freq}')
        
        # Show summer months
        summer_months = checklist.get('do_summer_months_covered', [])
        if summer_months:
            doc.add_paragraph(f'   Summer months sampled: {", ".join(summer_months)}')
        else:
            doc.add_paragraph('   Summer months sampled: Not specified (CONCERN)')
        
        # Nutrients Section
        doc.add_heading('Nutrients', 3)
        
        # Phosphorus
        p = doc.add_paragraph()
        p.add_run('Phosphorus (SRP/Orthophosphate):').bold = True
        
        nutrient_items = [
            ('srp_measured', 'SRP measured?'),
            ('srp_below_oxycline', 'SRP measured below oxycline (in hypoxic water)?'),
            ('srp_multi_year_available', 'Multi-year SRP data?'),
            ('biomass_potential_calculated', 'Phytoplankton Biomass Potential calculated?'),
        ]
        
        for key, label in nutrient_items:
            value = checklist.get(key, False)
            symbol = '✓' if value else '✗'
            doc.add_paragraph(f'{symbol} {label}', style='List Bullet')
        
        # Nitrogen
        p = doc.add_paragraph()
        p.add_run('Nitrogen (Ammonia):').bold = True
        
        ammonia_items = [
            ('ammonia_measured', 'Ammonia measured?'),
            ('ammonia_below_oxycline', 'Ammonia measured below oxycline?'),
            ('ammonia_multi_year_available', 'Multi-year ammonia data?'),
        ]
        
        for key, label in ammonia_items:
            value = checklist.get(key, False)
            symbol = '✓' if value else '✗'
            doc.add_paragraph(f'{symbol} {label}', style='List Bullet')
        
        # Phytoplankton Section
        doc.add_heading('Phytoplankton', 3)
        
        phyto_items = [
            ('phytoplankton_taxonomy_detailed', 'Detailed taxonomy (species identification)?'),
            ('phytoplankton_biovolume_by_taxa', 'Biovolume by taxa?'),
            ('phytoplankton_cell_count_by_taxa', 'Cell count by taxa?'),
            ('phytoplankton_multi_year_available', 'Multi-year phytoplankton data?'),
        ]
        
        for key, label in phyto_items:
            value = checklist.get(key, False)
            symbol = '✓' if value else '✗'
            doc.add_paragraph(f'{symbol} {label}', style='List Bullet')
        
        # Calculated Values Section (if available)
        has_calculations = (checklist.get('calculated_hypoxic_volume_m3') or 
                          checklist.get('calculated_biomass_potential_tonnes'))
        
        if has_calculations:
            doc.add_heading('Our Calculations', 3)
            
            if checklist.get('calculated_hypoxic_volume_m3'):
                vol = checklist['calculated_hypoxic_volume_m3']
                pct = checklist.get('calculated_hypoxic_percentage', 0)
                p = doc.add_paragraph()
                p.add_run('Hypoxic Water Volume: ').bold = True
                p.add_run(f'{vol:,.0f} m³')
                if pct:
                    p.add_run(f' ({pct:.1f}% of lake)')
            
            if checklist.get('calculated_available_p_kg'):
                p_kg = checklist['calculated_available_p_kg']
                p = doc.add_paragraph()
                p.add_run('Available Phosphorus: ').bold = True
                p.add_run(f'{p_kg:.1f} kg')
            
            if checklist.get('calculated_biomass_potential_tonnes'):
                tonnes = checklist['calculated_biomass_potential_tonnes']
                p = doc.add_paragraph()
                p.add_run('Phytoplankton Biomass Potential: ').bold = True
                p.add_run(f'{tonnes:.1f} metric tonnes')
                doc.add_paragraph('   (Potential algal bloom capacity from available phosphorus)')
        
        # Criticisms Section
        criticisms = checklist.get('criticisms', [])
        if criticisms:
            doc.add_heading('Analysis Criticisms', 3)
            doc.add_paragraph('Issues identified with the report\'s approach:')
            for criticism in criticisms[:5]:
                doc.add_paragraph(f'• {criticism}', style='List Bullet')
        
        doc.add_paragraph()  # Spacing
    
    def _add_comprehensive_analysis(self, doc: Document, evaluation: Dict):
        """Add comprehensive data analysis section with extracted values and calculations"""
        
        comp_analysis = evaluation.get('comprehensive_analysis', {})
        extracted_data = evaluation.get('extracted_data_values', comp_analysis.get('extracted_data', {}))
        
        # First add background and specific values (per client examples)
        self._add_lake_background(doc, evaluation)
        self._add_specific_values(doc, evaluation)
        
        # Add comprehensive checklist (NEW - per client feedback)
        self._add_data_quality_checklist(doc, evaluation)
        
        if not comp_analysis and not extracted_data:
            return
        
        doc.add_paragraph()
        doc.add_heading('Data Analysis Deep Dive', 2)
        
        # Extracted Data Values section
        if extracted_data and not extracted_data.get('error'):
            doc.add_heading('Key Data Values Extracted', 3)
            
            # DO data
            if extracted_data.get('minimum_do_value'):
                p = doc.add_paragraph()
                p.add_run('Dissolved Oxygen: ').bold = True
                min_do = extracted_data['minimum_do_value']
                p.add_run(f'Minimum DO = {min_do} mg/L')
                if min_do < 2:
                    p.add_run(' (SEVERE HYPOXIA)')
                elif min_do < 4:
                    p.add_run(' (MODERATE HYPOXIA)')
            
            if extracted_data.get('minimum_do_depth'):
                doc.add_paragraph(f"    Found at depth: {extracted_data['minimum_do_depth']}m")
            
            if extracted_data.get('oxycline_depth'):
                doc.add_paragraph(f"    Oxycline depth: {extracted_data['oxycline_depth']}m")
            
            # DO profile description
            if extracted_data.get('do_profile_description'):
                doc.add_paragraph(f"    Profile: {extracted_data['do_profile_description']}")
            
            # SRP data
            if extracted_data.get('srp_in_hypolimnion'):
                p = doc.add_paragraph()
                p.add_run('Phosphorus (SRP): ').bold = True
                p.add_run(f"Hypolimnion SRP = {extracted_data['srp_in_hypolimnion']} mg/L")
            
            # Nutrient profile description
            if extracted_data.get('nutrient_profile_description'):
                doc.add_paragraph(f"    Profile: {extracted_data['nutrient_profile_description']}")
            
            # Lake morphometry
            if extracted_data.get('max_depth') or extracted_data.get('lake_volume'):
                p = doc.add_paragraph()
                p.add_run('Lake Morphometry: ').bold = True
                morph_items = []
                if extracted_data.get('max_depth'):
                    morph_items.append(f"Max depth = {extracted_data['max_depth']}m")
                if extracted_data.get('lake_volume'):
                    morph_items.append(f"Volume = {extracted_data['lake_volume']:,.0f} m³")
                if extracted_data.get('lake_surface_area'):
                    morph_items.append(f"Area = {extracted_data['lake_surface_area']} ha")
                p.add_run(', '.join(morph_items))
            
            # Multi-year data
            years = extracted_data.get('years_with_data', [])
            if len(years) >= 2:
                p = doc.add_paragraph()
                p.add_run('Multi-Year Data: ').bold = True
                p.add_run(f"Data available for {len(years)} years ({min(years)}-{max(years)})")
        
        # Calculations Assessment section
        calc_assessment = comp_analysis.get('calculations_assessment', {})
        if calc_assessment and not calc_assessment.get('error'):
            doc.add_heading('Calculations Assessment', 3)
            
            # What's calculated vs what could be
            if calc_assessment.get('hypoxic_volume_calculated'):
                doc.add_paragraph('✓ Hypoxic volume is calculated', style='List Bullet')
            elif calc_assessment.get('hypoxic_volume_could_be_calculated'):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run('⚠ Hypoxic volume COULD be calculated but IS NOT').bold = True
                p.add_run(' - Data exists but isn\'t being used!')
            
            if calc_assessment.get('biomass_potential_could_be_calculated'):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run('⚠ Phytoplankton Biomass Potential could be calculated').bold = True
                p.add_run(' - Use formula: Hypoxic Volume × SRP × 100 = tonnes of potential bloom')
            
            # What's missing for calculations
            missing_hypoxic = calc_assessment.get('missing_for_hypoxic_volume', [])
            if missing_hypoxic:
                doc.add_paragraph(f"To calculate hypoxic volume, need: {', '.join(missing_hypoxic)}")
            
            # Underutilized data
            underutilized = calc_assessment.get('underutilized_data', [])
            if underutilized:
                doc.add_heading('Data Available but NOT Fully Utilized', 3)
                for item in underutilized[:5]:
                    doc.add_paragraph(f'⚠ {item}', style='List Bullet')
        
        # Trend Analysis section
        trend_analysis = comp_analysis.get('trend_analysis', {})
        if trend_analysis and trend_analysis.get('has_multi_year_data'):
            doc.add_heading('Multi-Year Trend Analysis', 3)
            
            years_analyzed = trend_analysis.get('years_analyzed', [])
            if years_analyzed:
                doc.add_paragraph(f"Analysis period: {min(years_analyzed)} - {max(years_analyzed)}")
            
            # Parameter trends
            do_trend = trend_analysis.get('do_trend', 'unknown')
            p_trend = trend_analysis.get('phosphorus_trend', 'unknown')
            cyano_trend = trend_analysis.get('cyanobacteria_trend', 'unknown')
            
            if do_trend != 'unknown':
                trend_icon = '↑' if do_trend == 'improving' else ('↓' if do_trend == 'declining' else '→')
                doc.add_paragraph(f"{trend_icon} Dissolved Oxygen: {do_trend.upper()}")
            
            if p_trend != 'unknown':
                trend_icon = '↓' if p_trend == 'improving' else ('↑' if p_trend == 'declining' else '→')
                doc.add_paragraph(f"{trend_icon} Phosphorus: {p_trend.upper()}")
            
            if cyano_trend != 'unknown':
                trend_icon = '↓' if cyano_trend == 'improving' else ('↑' if cyano_trend == 'declining' else '→')
                doc.add_paragraph(f"{trend_icon} Cyanobacteria: {cyano_trend.upper()}")
            
            # Key year-over-year changes
            key_changes = trend_analysis.get('key_changes', [])
            if key_changes:
                doc.add_paragraph('Key year-over-year changes:')
                for change in key_changes[:3]:
                    doc.add_paragraph(f'• {change}', style='List Bullet')
        
        # Critical findings
        critical_findings = comp_analysis.get('critical_findings', [])
        if critical_findings:
            doc.add_heading('Critical Findings', 3)
            for finding in critical_findings[:5]:
                if 'CRITICAL' in finding:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f'🔴 {finding}').bold = True
                elif 'WARNING' in finding or 'CONCERNING' in finding:
                    doc.add_paragraph(f'🟡 {finding}', style='List Bullet')
                elif 'POSITIVE' in finding:
                    doc.add_paragraph(f'🟢 {finding}', style='List Bullet')
                else:
                    doc.add_paragraph(f'• {finding}', style='List Bullet')
    
    def _add_priority_actions(self, doc: Document, evaluation: Dict):
        """Add top 3-5 priority actions"""
        
        doc.add_heading('Your Top Priority Actions', 1)
        
        # Determine priorities based on evaluation
        priorities = []
        
        # Check for missing critical parameters
        missing_params = evaluation.get('critical_parameters', {}).get('missing', [])
        if missing_params:
            # Extract parameter names from dict objects
            param_names = []
            for param in missing_params[:2]:
                if isinstance(param, dict):
                    param_names.append(param.get('name', str(param)))
                else:
                    param_names.append(str(param))
            
            if param_names:
                priorities.append({
                    'action': f"Start measuring: {', '.join(param_names)}",
                    'why': "These are essential for understanding lake health",
                    'urgency': 'HIGH'
                })
        
        # Check for missing calculations
        missing_calcs = evaluation.get('critical_calculations', {}).get('missing', [])
        if missing_calcs:
            # Extract calculation name from dict object
            calc_name = missing_calcs[0]
            if isinstance(calc_name, dict):
                calc_name = calc_name.get('name', str(calc_name))
            else:
                calc_name = str(calc_name)
            
            priorities.append({
                'action': f"Calculate: {calc_name}",
                'why': "Critical for quantifying the problem extent",
                'urgency': 'HIGH'
            })
        
        # Check for problematic parameters
        problematic = evaluation.get('problematic_parameters', {}).get('found', [])
        if len(problematic) > 3:
            priorities.append({
                'action': "Remove non-actionable parameters from monitoring",
                'why': f"You're wasting resources on {len(problematic)} metrics that cannot be actioned upon",
                'urgency': 'MEDIUM'
            })
        
        # Add general improvements
        if evaluation.get('compliance_percentage', 0) < 70:
            priorities.append({
                'action': "Focus on root causes, not symptoms",
                'why': "Current approach treats symptoms without addressing underlying issues",
                'urgency': 'HIGH'
            })
        
        # Display top 3 priorities
        for i, priority in enumerate(priorities[:3], 1):
            para = doc.add_paragraph()
            para.add_run(f"{i}. {priority['action']}\n").bold = True
            para.add_run(f"   Why: {priority['why']}")
            
        doc.add_paragraph()  # Spacing
    
    def _add_score_summary(self, doc: Document, evaluation: Dict):
        """Add visual score summary"""
        
        doc.add_heading('Performance Summary', 1)
        
        # Create simple table for scores
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Headers
        table.cell(0, 0).text = 'Category'
        table.cell(0, 1).text = 'Status'
        
        # Critical parameters
        critical = evaluation.get('critical_parameters', {})
        found = len(critical.get('found', []))
        missing = len(critical.get('missing', []))
        table.cell(1, 0).text = 'Critical Parameters'
        table.cell(1, 1).text = f'{found} found, {missing} missing'
        
        # Calculations
        calcs = evaluation.get('critical_calculations', {})
        performed = len(calcs.get('performed', []))
        not_done = len(calcs.get('missing', []))
        table.cell(2, 0).text = 'Key Calculations'
        table.cell(2, 1).text = f'{performed} performed, {not_done} missing'
        
        # Overall
        score = evaluation.get('compliance_percentage', 0)
        table.cell(3, 0).text = 'Overall Score'
        table.cell(3, 1).text = f'{score:.0f}%'
        
        doc.add_paragraph()  # Spacing
    
    def _add_next_steps(self, doc: Document, doc_data: Dict, evaluation: Dict):
        """Add clear next steps section"""
        
        doc.add_heading('Your Next Steps', 1)
        
        # Get document type from evaluation (hybrid analysis)
        doc_type = evaluation.get('document_type', doc_data.get('document_type', 'hybrid'))
        type_breakdown = evaluation.get('type_breakdown', {})
        plan_pct = type_breakdown.get('plan_percentage', 50)
        report_pct = type_breakdown.get('report_percentage', 50)
        
        # Determine steps based on document composition
        if plan_pct > 70:
            # Primarily a Lake Management Plan
            steps = [
                "Review the priority actions above",
                "Add missing critical parameters to your monitoring plan",
                "Set up monthly data collection for key metrics",
                "Schedule quarterly review of progress",
                "Contact us for implementation support"
            ]
        elif report_pct > 70:
            # Primarily a Lake Report
            steps = [
                "Review which parameters provide actionable insights",
                "Calculate hypoxic water volume and sediment area",
                "Track trends monthly, not just annually",
                "Focus interventions on reducing hypoxia",
                "Request a consultation for targeted solutions"
            ]
        else:
            # Hybrid document - combined steps
            steps = [
                "Review both historical data quality AND future plan alignment",
                "Ensure proposed interventions address problems identified in data",
                "Add missing critical measurements to your monitoring plan",
                "Calculate hypoxic volume to quantify the root cause",
                "Schedule a consultation to bridge the data-to-action gap"
            ]
        
        for step in steps[:4]:
            doc.add_paragraph(step, style='List Number')
        
        doc.add_paragraph()  # Spacing
    
    def _add_overview_and_actions(self, doc: Document, evaluation: Dict):
        """Add OVERVIEW and ACTIONS sections per client example format"""
        
        comp_analysis = evaluation.get('comprehensive_analysis', {})
        checklist = comp_analysis.get('data_quality_checklist', {})
        lake_bg = comp_analysis.get('lake_background', {})
        
        # OVERVIEW Section
        doc.add_heading('OVERVIEW', 1)
        
        # Build overview based on analysis
        overview_points = []
        
        # Lake status
        lake_name = lake_bg.get('lake_name', 'This lake')
        if comp_analysis.get('trend_analysis', {}).get('water_quality_trend') == 'declining':
            overview_points.append(f"{lake_name} continues to deteriorate based on the data.")
        elif comp_analysis.get('trend_analysis', {}).get('water_quality_trend') == 'improving':
            overview_points.append(f"{lake_name} shows signs of improvement based on the data.")
        
        # Hypoxia status
        if checklist.get('calculated_hypoxic_percentage'):
            pct = checklist['calculated_hypoxic_percentage']
            if pct > 30:
                overview_points.append(f"The lake suffers from extensive hypoxia ({pct:.0f}% hypoxic), providing conditions that favor cyanobacteria.")
        elif not checklist.get('hypoxic_volume_calculated'):
            overview_points.append("Hypoxic water volume has not been calculated - this is a critical gap in understanding.")
        
        # Phytoplankton status
        if checklist.get('phytoplankton_taxonomy_detailed'):
            overview_points.append("Phytoplankton data includes detailed taxonomy for informed management.")
        else:
            overview_points.append("Phytoplankton analysis lacks detailed taxonomy needed for effective management.")
        
        # Current systems issues
        system_issues = lake_bg.get('system_issues', [])
        if system_issues:
            for issue in system_issues[:2]:
                overview_points.append(issue)
        
        # Display overview
        for point in overview_points[:5]:
            doc.add_paragraph(f'• {point}', style='List Bullet')
        
        if not overview_points:
            doc.add_paragraph("Insufficient data extracted for comprehensive overview. More detailed document analysis needed.")
        
        # RECOMMENDED ACTIONS Section
        doc.add_heading('RECOMMENDED ACTIONS', 1)
        
        actions = []
        
        # Action based on hypoxia
        if not checklist.get('hypoxic_volume_calculated') and checklist.get('hypsographic_table_available'):
            actions.append("Calculate hypoxic water volume using the available hypsographic table and DO profiles.")
        
        # Action based on DO
        if not checklist.get('do_measured_at_intervals_to_bottom'):
            actions.append("Implement DO measurement at regular intervals all the way to the lake bottom.")
        
        # Action based on nutrients
        if not checklist.get('srp_below_oxycline'):
            actions.append("Measure orthophosphate (SRP) below the oxycline to understand nutrient availability for cyanobacteria.")
        
        if not checklist.get('ammonia_below_oxycline'):
            actions.append("Measure ammonia below the oxycline - this is a key driver of cyanobacteria proliferation.")
        
        # Action based on phytoplankton
        if not checklist.get('phytoplankton_taxonomy_detailed'):
            actions.append("Improve phytoplankton analysis with detailed taxonomy to distinguish beneficial algae from cyanobacteria.")
        
        # Action for biomass potential
        if not checklist.get('biomass_potential_calculated') and checklist.get('srp_below_oxycline'):
            actions.append("Calculate Phytoplankton Biomass Potential using SRP and hypoxic volume data.")
        
        # Generic actions if we don't have specific ones
        if not actions:
            actions = [
                "Focus monitoring resources on critical parameters (DO, SRP, ammonia, phytoplankton)",
                "Calculate hypoxic water volume to quantify the extent of the problem",
                "Track data monthly during summer months for better trend analysis"
            ]
        
        for action in actions[:5]:
            doc.add_paragraph(f'• {action}', style='List Bullet')
        
        doc.add_paragraph()  # Spacing
    
    def _add_contact_footer(self, doc: Document):
        """Add contact information footer"""
        
        doc.add_page_break()
        
        # Call to action (softened language per client feedback)
        cta = doc.add_heading('Ready to Improve Your Lake Management?', 2)
        cta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run("Reallocate your resources to the measurements that matter most.\n")
        para.add_run("Get expert guidance to create an effective ACTION plan.\n\n")
        
        # Contact box
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.add_run("Schedule Your Consultation\n").bold = True
        contact.add_run("📞 1-800-LAKE-FIX\n")
        contact.add_run("✉️ action@reporttoreveal.com\n")
        contact.add_run("🌐 www.reporttoreveal.com\n\n")
        
        # Tagline
        tagline = doc.add_paragraph()
        tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tagline.add_run("Report to Reveal, Not Conceal™").italic = True
    
    def generate_trend_summary(
        self,
        multi_year_data: List[Dict],
        output_path: Optional[Path] = None
    ) -> str:
        """
        Generate summary for multi-year trend analysis
        
        Args:
            multi_year_data: List of evaluations for multiple years
            output_path: Where to save the trend summary
            
        Returns:
            Path to generated trend summary
        """
        try:
            doc = Document()
            
            # Set margins
            for section in doc.sections:
                section.top_margin = Inches(0.8)
                section.bottom_margin = Inches(0.8)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Title
            title = doc.add_heading('Lake Assessment - Trend Analysis', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Period covered
            years = [data.get('year', 'Unknown') for data in multi_year_data]
            period = f"{min(years)} - {max(years)}" if years else "Multi-Year"
            
            info = doc.add_paragraph()
            info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            info.add_run(f'Analysis Period: {period}\n').bold = True
            info.add_run(f'Reports Analyzed: {len(multi_year_data)}')
            
            doc.add_paragraph()
            
            # Trend summary
            doc.add_heading('Key Trends', 1)
            
            # Calculate trend direction
            if len(multi_year_data) >= 2:
                first_score = multi_year_data[0].get('compliance_percentage', 0)
                last_score = multi_year_data[-1].get('compliance_percentage', 0)
                trend = last_score - first_score
                
                if trend > 5:
                    trend_text = f"✓ Improving: Score increased by {trend:.0f}%"
                    trend_color = self.accent_color
                elif trend < -5:
                    trend_text = f"✗ Declining: Score decreased by {abs(trend):.0f}%"
                    trend_color = RGBColor(220, 53, 69)
                else:
                    trend_text = f"→ Stable: Score changed by {trend:.0f}%"
                    trend_color = self.warning_color
                
                para = doc.add_paragraph()
                trend_run = para.add_run(trend_text)
                trend_run.font.size = Pt(12)
                trend_run.bold = True
            
            # Year-by-year summary table
            doc.add_heading('Year-by-Year Performance', 2)
            
            table = doc.add_table(rows=len(multi_year_data)+1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            # Headers
            table.cell(0, 0).text = 'Year'
            table.cell(0, 1).text = 'Score'
            table.cell(0, 2).text = 'Key Finding'
            
            # Data rows
            for i, data in enumerate(multi_year_data, 1):
                table.cell(i, 0).text = str(data.get('year', 'Year ' + str(i)))
                table.cell(i, 1).text = f"{data.get('compliance_percentage', 0):.0f}%"
                
                # Determine key finding
                missing = len(data.get('critical_parameters', {}).get('missing', []))
                if missing > 3:
                    finding = f"{missing} critical parameters missing"
                elif data.get('compliance_percentage', 0) < 50:
                    finding = "Needs significant improvement"
                else:
                    finding = "Acceptable performance"
                table.cell(i, 2).text = finding
            
            # Recommendations
            doc.add_heading('Recommendations Based on Trends', 1)
            
            if trend < -5:
                recommendations = [
                    "Urgent intervention needed to reverse declining trend",
                    "Review and update monitoring protocols immediately",
                    "Consider external consultation for rapid improvement",
                    "Implement monthly progress reviews"
                ]
            elif trend > 5:
                recommendations = [
                    "Continue current improvement trajectory",
                    "Document successful interventions for replication",
                    "Set higher targets for next period",
                    "Share success strategies with stakeholders"
                ]
            else:
                recommendations = [
                    "Break out of stagnation with new approaches",
                    "Review why current efforts aren't improving outcomes",
                    "Consider adopting best practices from successful lakes",
                    "Set specific, measurable improvement goals"
                ]
            
            for rec in recommendations:
                doc.add_paragraph(f'{rec}', style='List Bullet')
            
            # Save
            if output_path is None:
                output_path = Path("results") / f"trend_summary_{datetime.now().strftime('%Y%m%d')}.docx"
            
            doc.save(str(output_path))
            logger.info(f"Trend summary generated: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Error generating trend summary: {e}")
            raise

# Global instance
summary_generator = SummaryReportGenerator()
