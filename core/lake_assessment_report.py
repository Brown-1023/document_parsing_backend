"""
Lake Assessment Report Generator
Creates comprehensive multi-year trend analysis reports
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os
from typing import Dict, List
import logging

logger = logging.getLogger(__name__)


class LakeAssessmentReportGenerator:
    """Generates Lake Assessment reports for multi-year trend analysis"""
    
    def __init__(self, output_dir: str = "results"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
    
    def generate_assessment_report(self, assessment_data: Dict) -> str:
        """
        Generate a Lake Assessment report
        
        Args:
            assessment_data: Assessment data for a single lake
            
        Returns:
            Path to generated report file
        """
        doc = Document()
        
        # Set up styles
        self._setup_styles(doc)
        
        # Add header
        self._add_header(doc, assessment_data)
        
        # Executive Summary
        self._add_executive_summary(doc, assessment_data)
        
        # Trend Analysis Results
        self._add_trend_analysis(doc, assessment_data)
        
        # Year-by-Year Comparison
        self._add_yearly_comparison(doc, assessment_data)
        
        # Key Findings
        self._add_key_findings(doc, assessment_data)
        
        # Recommendations
        self._add_recommendations(doc, assessment_data)
        
        # Technical Appendix
        self._add_technical_appendix(doc, assessment_data)
        
        # Save document
        lake_name = assessment_data.get('lake_name', 'Unknown Lake')
        safe_name = "".join(c for c in lake_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
        filename = f"{safe_name}_Lake_Assessment_{datetime.now().strftime('%Y%m%d')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        
        doc.save(filepath)
        logger.info(f"Lake Assessment report generated: {filepath}")
        
        return filepath
    
    def _setup_styles(self, doc: Document):
        """Set up document styles"""
        styles = doc.styles
        
        # Modify Normal style
        style = styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Modify Heading styles
        for i in range(1, 4):
            style = styles[f'Heading {i}']
            style.font.name = 'Calibri'
            style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    
    def _add_header(self, doc: Document, assessment_data: Dict):
        """Add document header"""
        # Title
        title = doc.add_heading('Lake Assessment Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle with lake name
        lake_name = assessment_data.get('lake_name', 'Unknown Lake')
        subtitle = doc.add_heading(f'Multi-Year Trend Analysis for {lake_name}', 2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Report metadata
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run(f"Analysis Period: {assessment_data.get('year_range', 'N/A')}\n").bold = True
        para.add_run(f"Reports Analyzed: {assessment_data.get('reports_analyzed', 0)}\n")
        para.add_run(f"Assessment Date: {datetime.now().strftime('%B %d, %Y')}\n")
        
        doc.add_paragraph()
        doc.add_paragraph("─" * 80)
        doc.add_paragraph()
    
    def _add_executive_summary(self, doc: Document, assessment_data: Dict):
        """Add executive summary"""
        doc.add_heading('Executive Summary', 1)
        
        trend_analysis = assessment_data.get('trend_analysis', {})
        trajectory = trend_analysis.get('overall_trajectory', 'Unknown')
        
        # Overall status box
        para = doc.add_paragraph()
        para.add_run('Overall Lake Trajectory: ').bold = True
        
        # Color-code the trajectory
        run = para.add_run(trajectory)
        run.bold = True
        if 'Degradation' in trajectory:
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        elif 'Improvement' in trajectory:
            run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        else:
            run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
        
        doc.add_paragraph()
        
        # Summary text
        years = trend_analysis.get('years', [])
        if len(years) >= 3:
            summary_text = (
                f"This assessment analyzes {assessment_data.get('reports_analyzed', 0)} reports "
                f"from {assessment_data.get('lake_name', 'the lake')} spanning {assessment_data.get('year_range', 'multiple years')}. "
                f"The analysis reveals trends in key water quality parameters, hypoxia development, "
                f"and overall lake health trajectory."
            )
            doc.add_paragraph(summary_text)
        else:
            doc.add_paragraph(
                "Insufficient multi-year data for comprehensive trend analysis. "
                "At least 3 years of data are required for meaningful trend assessment."
            )
        
        doc.add_paragraph()
    
    def _add_trend_analysis(self, doc: Document, assessment_data: Dict):
        """Add detailed trend analysis section"""
        doc.add_heading('Parameter Trend Analysis', 1)
        
        trend_analysis = assessment_data.get('trend_analysis', {})
        parameters = trend_analysis.get('parameters', {})
        
        if not parameters:
            doc.add_paragraph('No parameter trends available for analysis.')
            return
        
        # Create a table for trend summary
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Parameter'
        header_cells[1].text = 'Trend Direction'
        header_cells[2].text = 'Change Rate'
        header_cells[3].text = 'First Year Value'
        header_cells[4].text = 'Last Year Value'
        
        # Make header bold
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
        
        # Add parameter rows
        param_display_names = {
            'dissolved_oxygen_min': 'Min Dissolved Oxygen',
            'hypoxic_volume': 'Hypoxic Water Volume',
            'hypoxic_percentage': 'Hypoxic Percentage',
            'orthophosphate_max': 'Max Orthophosphate',
            'ammonia_max': 'Max Ammonia',
            'cyanobacteria_percentage': 'Cyanobacteria %',
            'compliance_score': 'Compliance Score'
        }
        
        for param, trend_data in parameters.items():
            if trend_data.get('trend') == 'analyzed':
                row_cells = table.add_row().cells
                row_cells[0].text = param_display_names.get(param, param)
                
                # Direction with arrow
                direction = trend_data.get('direction', 'unknown')
                if direction == 'increasing':
                    row_cells[1].text = '↑ Increasing'
                elif direction == 'decreasing':
                    row_cells[1].text = '↓ Decreasing'
                elif direction == 'stable':
                    row_cells[1].text = '→ Stable'
                else:
                    row_cells[1].text = 'No clear trend'
                
                # Change rate
                if trend_data.get('change_rate') is not None:
                    row_cells[2].text = f"{trend_data['change_rate']:.1f}%"
                else:
                    row_cells[2].text = 'N/A'
                
                # Values
                if trend_data.get('first_value') is not None:
                    row_cells[3].text = f"{trend_data['first_value']:.2f}"
                else:
                    row_cells[3].text = 'N/A'
                    
                if trend_data.get('last_value') is not None:
                    row_cells[4].text = f"{trend_data['last_value']:.2f}"
                else:
                    row_cells[4].text = 'N/A'
        
        doc.add_paragraph()
        
        # Add interpretation
        doc.add_heading('Trend Interpretation', 2)
        
        # Hypoxia trends
        if 'hypoxic_volume' in parameters or 'hypoxic_percentage' in parameters:
            para = doc.add_paragraph()
            para.add_run('Hypoxia Trends: ').bold = True
            
            if 'hypoxic_volume' in parameters:
                trend = parameters['hypoxic_volume']
                if trend.get('direction') == 'increasing':
                    para.add_run(
                        'Hypoxic water volume is expanding, indicating deteriorating oxygen conditions. '
                        'This trend increases HAB risk and requires immediate intervention.'
                    )
                elif trend.get('direction') == 'decreasing':
                    para.add_run(
                        'Hypoxic water volume is reducing, suggesting improving oxygen conditions. '
                        'Continue current management strategies.'
                    )
        
        # Nutrient trends
        if 'orthophosphate_max' in parameters or 'ammonia_max' in parameters:
            para = doc.add_paragraph()
            para.add_run('Nutrient Trends: ').bold = True
            
            nutrients_increasing = False
            for param in ['orthophosphate_max', 'ammonia_max']:
                if param in parameters:
                    if parameters[param].get('direction') == 'increasing':
                        nutrients_increasing = True
                        break
            
            if nutrients_increasing:
                para.add_run(
                    'Nutrient levels are increasing, providing more fuel for algal growth. '
                    'This trend requires attention to watershed management and nutrient sources.'
                )
            else:
                para.add_run(
                    'Nutrient levels are stable or decreasing, indicating effective nutrient management.'
                )
        
        doc.add_paragraph()
    
    def _add_yearly_comparison(self, doc: Document, assessment_data: Dict):
        """Add year-by-year comparison"""
        doc.add_heading('Year-by-Year Comparison', 1)
        
        individual_reports = assessment_data.get('individual_reports', [])
        
        if not individual_reports:
            doc.add_paragraph('No individual report data available.')
            return
        
        # Create comparison table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Year'
        header_cells[1].text = 'Compliance Score'
        header_cells[2].text = 'Risk Level'
        header_cells[3].text = 'Key Issues'
        
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
        
        # Add yearly data
        for report in individual_reports:
            year = report.get('extracted_year', 'Unknown')
            compliance = report.get('compliance_evaluation', {})
            
            row_cells = table.add_row().cells
            row_cells[0].text = str(year)
            
            score = compliance.get('overall_score')
            if score:
                row_cells[1].text = f"{score:.1f}%"
            else:
                row_cells[1].text = 'N/A'
            
            # Risk level from lake condition assessment
            lake_assessment = compliance.get('lake_condition_assessment', {})
            risk = lake_assessment.get('risk_level', 'Unknown')
            row_cells[2].text = risk
            
            # Key issues
            issues = []
            if compliance.get('critical_parameters', {}).get('missing'):
                issues.append(f"{len(compliance['critical_parameters']['missing'])} missing critical params")
            if compliance.get('problematic_parameters', {}).get('found'):
                issues.append(f"{len(compliance['problematic_parameters']['found'])} problematic params")
            
            row_cells[3].text = ', '.join(issues) if issues else 'None identified'
        
        doc.add_paragraph()
    
    def _add_key_findings(self, doc: Document, assessment_data: Dict):
        """Add key findings section"""
        doc.add_heading('Key Findings', 1)
        
        trend_analysis = assessment_data.get('trend_analysis', {})
        findings = trend_analysis.get('key_findings', [])
        
        if findings:
            for finding in findings:
                doc.add_paragraph(finding, style='List Bullet')
        else:
            doc.add_paragraph('No significant findings identified.')
        
        doc.add_paragraph()
    
    def _add_recommendations(self, doc: Document, assessment_data: Dict):
        """Add recommendations section"""
        doc.add_heading('Recommendations Based on Trends', 1)
        
        trend_analysis = assessment_data.get('trend_analysis', {})
        recommendations = trend_analysis.get('recommendations', [])
        
        if recommendations:
            for i, rec in enumerate(recommendations, 1):
                para = doc.add_paragraph()
                para.add_run(f"{i}. ").bold = True
                para.add_run(rec)
        else:
            doc.add_paragraph('Continue current monitoring and management practices.')
        
        # Add monitoring recommendations
        doc.add_heading('Monitoring Recommendations', 2)
        
        doc.add_paragraph(
            "Based on the trend analysis, we recommend:",
            style='Normal'
        )
        
        monitoring_recs = [
            "Continue monthly monitoring of all critical parameters",
            "Ensure consistent measurement locations and methods across years",
            "Document any changes in watershed conditions or management practices",
            "Perform annual trend analysis to track progress"
        ]
        
        for rec in monitoring_recs:
            doc.add_paragraph(rec, style='List Bullet')
        
        doc.add_paragraph()
    
    def _add_technical_appendix(self, doc: Document, assessment_data: Dict):
        """Add technical appendix with methodology"""
        doc.add_page_break()
        doc.add_heading('Technical Appendix', 1)
        
        doc.add_heading('Methodology', 2)
        doc.add_paragraph(
            "This Lake Assessment uses linear regression analysis to identify statistically significant trends "
            "in water quality parameters over time. Trends are classified as:"
        )
        
        methodology_items = [
            "Increasing: Positive slope with p-value < 0.05",
            "Decreasing: Negative slope with p-value < 0.05",
            "Stable: Slope near zero (|slope| < 0.01) with p-value < 0.05",
            "No clear trend: p-value ≥ 0.05 (not statistically significant)"
        ]
        
        for item in methodology_items:
            doc.add_paragraph(item, style='List Bullet')
        
        doc.add_heading('Data Sources', 2)
        doc.add_paragraph(
            f"This assessment is based on {assessment_data.get('reports_analyzed', 0)} reports "
            f"submitted for {assessment_data.get('lake_name', 'the lake')}:"
        )
        
        # List analyzed reports
        for report in assessment_data.get('individual_reports', []):
            year = report.get('extracted_year', 'Unknown')
            filename = report.get('filename', 'Unknown file')
            doc.add_paragraph(f"Year {year}: {filename}", style='List Bullet')
        
        doc.add_heading('Limitations', 2)
        doc.add_paragraph(
            "This trend analysis is subject to the following limitations:"
        )
        
        limitations = [
            "Trends assume consistent measurement methods across years",
            "Missing data points may affect trend accuracy",
            "Natural variability and weather patterns are not accounted for",
            "Minimum of 3 data points required for meaningful trends"
        ]
        
        for limitation in limitations:
            doc.add_paragraph(limitation, style='List Bullet')
        
        # Footer
        doc.add_paragraph()
        doc.add_paragraph("─" * 80)
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("Lake Assessment Report - ").italic = True
        footer.add_run("Report to Reveal, Not Conceal™").bold = True
        footer.add_run("\nGenerated by Lake Management Analysis System").italic = True
