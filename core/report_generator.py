"""
Report generator for creating Word documents with analysis results
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from typing import Dict, Any, List, Tuple, Optional
import json
from pathlib import Path
from datetime import datetime
import logging
from .summary_generator import summary_generator

logger = logging.getLogger(__name__)

class ReportGenerator:
    """Generate professional Word document reports"""
    
    def __init__(self):
        self.brand_colors = {
            'primary': RGBColor(0, 51, 102),      # Navy blue
            'secondary': RGBColor(0, 128, 128),   # Teal
            'success': RGBColor(0, 128, 0),       # Green
            'warning': RGBColor(255, 140, 0),     # Orange
            'danger': RGBColor(220, 20, 60),      # Red
            'info': RGBColor(70, 130, 180)        # Steel blue
        }
    
    async def generate_report(
        self, 
        doc_data: Dict[str, Any],
        compliance_eval: Dict[str, Any],
        ai_insights: Dict[str, Any] = None,
        output_path: str = None,
        generate_summary: bool = True
    ) -> Tuple[str, Optional[str]]:
        """
        Generate comprehensive Word document report and optional summary
        
        Returns tuple of (detailed_report_path, summary_report_path)
        """
        doc = Document()
        self._setup_styles(doc)
        
        # Add header
        self._add_header(doc, doc_data)
        
        # Executive Summary
        self._add_executive_summary(doc, compliance_eval, ai_insights)
        
        # Hybrid Analysis Section (if available)
        if compliance_eval.get('hybrid_analysis') or compliance_eval.get('hybrid_summary'):
            self._add_hybrid_analysis_section(doc, compliance_eval, ai_insights)
        
        # Compliance Score Card
        self._add_score_card(doc, compliance_eval)
        
        # Individual Parameter Comments (Brief requirement)
        self._add_individual_parameter_comments(doc, compliance_eval)
        
        # Lake Assessment (Brief requirement)
        self._add_lake_assessment(doc, compliance_eval)
        
        # Parameter Analysis
        self._add_parameter_analysis(doc, compliance_eval)
        
        # Calculations Analysis
        self._add_calculations_analysis(doc, compliance_eval)
        
        # Phytoplankton Management Assessment (per client feedback)
        self._add_phytoplankton_management(doc, compliance_eval)
        
        # AI Insights (if available)
        if ai_insights and ai_insights.get("ai_analysis_available"):
            self._add_ai_insights(doc, ai_insights)
        
        # Recommendations
        self._add_recommendations(doc, compliance_eval)
        
        # Call to Action
        self._add_call_to_action(doc)
        
        # Appendix
        self._add_appendix(doc, doc_data, compliance_eval)
        
        # Save document
        if not output_path:
            filename = f"analysis_{doc_data.get('filename', 'report').replace('.pdf', '')}.docx"
            output_path = Path("results") / filename
        
        Path(output_path).parent.mkdir(exist_ok=True)
        doc.save(output_path)
        logger.info(f"Report generated: {output_path}")
        
        # Generate summary report if requested
        summary_path = None
        if generate_summary:
            try:
                # Create summary output path
                output_dir = Path(output_path).parent
                base_name = Path(output_path).stem
                summary_output = output_dir / f"{base_name}_summary.docx"
                
                # Generate the summary
                summary_path = summary_generator.generate_summary(
                    doc_data=doc_data,
                    evaluation=compliance_eval,
                    output_path=summary_output
                )
                logger.info(f"Summary generated: {summary_path}")
            except Exception as e:
                logger.error(f"Failed to generate summary: {e}")
                # Continue even if summary fails
        
        return str(output_path), summary_path
    
    def _setup_styles(self, doc: Document):
        """Setup custom styles for the document"""
        styles = doc.styles
        
        # Create custom heading styles
        if 'Custom Heading 1' not in styles:
            heading1 = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            heading1.font.name = 'Calibri'
            heading1.font.size = Pt(16)
            heading1.font.bold = True
            heading1.font.color.rgb = self.brand_colors['primary']
    
    def _add_header(self, doc: Document, doc_data: Dict):
        """Add report header with branding"""
        # Title
        title = doc.add_heading('Lake Management Document Analysis', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle with tagline
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('Report to Reveal, Not Conceal™')
        run.font.size = Pt(14)
        run.font.italic = True
        run.font.color.rgb = self.brand_colors['secondary']
        
        doc.add_paragraph()
        
        # Document type indicator - shows type based on percentages
        doc_type = doc_data.get('document_type', 'hybrid')
        type_breakdown = doc_data.get('type_breakdown', {})
        plan_pct = type_breakdown.get('plan_percentage', 50)
        report_pct = type_breakdown.get('report_percentage', 50)
        
        type_p = doc.add_paragraph()
        type_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if report_pct > 70:
            # Primarily a report
            type_run = type_p.add_run("LAKE REPORT ANALYSIS")
            type_run.bold = True
            type_run.font.size = Pt(12)
            type_run.font.color.rgb = RGBColor(0, 76, 153)  # Blue
            
            detail_note = doc.add_paragraph()
            detail_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
            note_run = detail_note.add_run(f"Primarily Retrospective Data ({report_pct}%)")
            note_run.font.size = Pt(10)
            note_run.font.italic = True
            note_run.font.color.rgb = RGBColor(100, 100, 100)
        elif plan_pct > 70:
            # Primarily a plan
            type_run = type_p.add_run("LAKE MANAGEMENT PLAN ANALYSIS")
            type_run.bold = True
            type_run.font.size = Pt(12)
            type_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
            
            detail_note = doc.add_paragraph()
            detail_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
            note_run = detail_note.add_run(f"Primarily Forward-Looking ({plan_pct}%)")
            note_run.font.size = Pt(10)
            note_run.font.italic = True
            note_run.font.color.rgb = RGBColor(100, 100, 100)
        else:
            # Hybrid
            type_run = type_p.add_run("HYBRID ANALYSIS")
            type_run.bold = True
            type_run.font.size = Pt(12)
            type_run.font.color.rgb = RGBColor(0, 128, 128)  # Teal
            
            detail_note = doc.add_paragraph()
            detail_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
            note_run = detail_note.add_run(f"Report: {report_pct}% | Plan: {plan_pct}%")
            note_run.font.size = Pt(10)
            note_run.font.italic = True
            note_run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Always show that both perspectives were analyzed
        both_note = doc.add_paragraph()
        both_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        both_run = both_note.add_run("(Analyzed from both retrospective and forward-looking perspectives)")
        both_run.font.size = Pt(9)
        both_run.font.italic = True
        both_run.font.color.rgb = RGBColor(120, 120, 120)
        
        # Document info
        info = doc.add_paragraph()
        info.add_run('Document Analyzed: ').bold = True
        info.add_run(doc_data.get('filename', 'Unknown'))
        doc.add_paragraph(f"Analysis Date: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph(f"Pages: {doc_data.get('page_count', 'N/A')}")
        
        doc.add_paragraph()
    
    def _add_executive_summary(self, doc: Document, compliance_eval: Dict, ai_insights: Dict):
        """Add executive summary section"""
        doc.add_heading('Executive Summary', 1)
        
        # Add "Our Thinking" comparison statement
        doc.add_paragraph(
            "This report has been analyzed against scientifically-backed best practices "
            "that focus on root causes rather than symptoms, as recommended by the GAO-22-104449 report "
            "demonstrating the inextricable link between hypoxia and Harmful Algal Blooms (HABs)."
        )
        doc.add_paragraph()
        
        # Overall assessment
        score = compliance_eval['compliance_percentage']
        level = compliance_eval['compliance_level'].value
        
        # Score paragraph with color coding
        p = doc.add_paragraph()
        p.add_run('Overall Compliance Score: ').bold = True
        
        score_run = p.add_run(f"{score:.1f}%")
        score_run.bold = True
        score_run.font.size = Pt(14)
        
        # Color based on score
        if score >= 80:
            score_run.font.color.rgb = self.brand_colors['success']
            assessment = "This report demonstrates excellent adherence to best practices."
        elif score >= 60:
            score_run.font.color.rgb = self.brand_colors['info']
            assessment = "This report shows good understanding but has room for improvement."
        elif score >= 40:
            score_run.font.color.rgb = self.brand_colors['warning']
            assessment = "This report has significant gaps in addressing root causes."
        else:
            score_run.font.color.rgb = self.brand_colors['danger']
            assessment = "This report focuses primarily on symptoms rather than root causes."
        
        p.add_run(f" ({level.upper()})")
        
        doc.add_paragraph(assessment)
        
        # Key findings
        doc.add_heading('Key Findings', 2)
        
        # Strengths
        if compliance_eval.get('strengths'):
            doc.add_paragraph('Strengths:')
            for strength in compliance_eval['strengths'][:3]:
                doc.add_paragraph(f"{strength}", style='List Bullet')
        
        # Weaknesses
        if compliance_eval.get('weaknesses'):
            doc.add_paragraph('Areas for Improvement:')
            for weakness in compliance_eval['weaknesses'][:3]:
                doc.add_paragraph(f"{weakness}", style='List Bullet')
        
        doc.add_paragraph()
    
    def _add_hybrid_analysis_section(self, doc: Document, compliance_eval: Dict, ai_insights: Dict):
        """Add hybrid analysis section showing both retrospective and forward-looking analysis"""
        doc.add_heading('Hybrid Document Analysis', 1)
        
        # Introduction
        doc.add_paragraph(
            "Most lake management documents contain BOTH historical data/reports AND future plans. "
            "Our analysis evaluates your document from both perspectives to provide comprehensive insights."
        )
        doc.add_paragraph()
        
        # Document Characteristics
        hybrid_analysis = compliance_eval.get('hybrid_analysis', {})
        hybrid_summary = compliance_eval.get('hybrid_summary', {})
        characteristics = hybrid_analysis.get('document_characteristics', {})
        
        if characteristics or hybrid_summary:
            doc.add_heading('Document Composition', 2)
            
            # Show percentages
            plan_pct = characteristics.get('plan_percentage', 50)
            report_pct = characteristics.get('report_percentage', 50)
            
            p = doc.add_paragraph()
            p.add_run('This document is: ').bold = True
            p.add_run(f"{report_pct}% retrospective data/reporting and {plan_pct}% forward-looking plans")
            
            dominant = characteristics.get('dominant_type', 'hybrid')
            doc.add_paragraph(f"Dominant Type: {dominant.upper()}")
            
            # Content breakdown
            if characteristics.get('has_historical_data'):
                doc.add_paragraph("✓ Contains historical data and measurements", style='List Bullet')
            if characteristics.get('has_future_plans'):
                doc.add_paragraph("✓ Contains future plans and recommendations", style='List Bullet')
            if characteristics.get('has_interventions'):
                doc.add_paragraph("✓ Mentions intervention strategies", style='List Bullet')
        
        # Retrospective Analysis (Report aspect)
        retrospective = hybrid_analysis.get('retrospective_analysis', {})
        if retrospective:
            doc.add_heading('Retrospective Analysis (Historical Data)', 2)
            
            data_quality = retrospective.get('data_quality', 'unknown')
            p = doc.add_paragraph()
            p.add_run('Data Quality: ').bold = True
            quality_run = p.add_run(data_quality.upper())
            if data_quality == 'good':
                quality_run.font.color.rgb = self.brand_colors['success']
            elif data_quality == 'fair':
                quality_run.font.color.rgb = self.brand_colors['warning']
            else:
                quality_run.font.color.rgb = self.brand_colors['danger']
            
            if retrospective.get('temporal_coverage'):
                doc.add_paragraph(f"Temporal Coverage: {retrospective['temporal_coverage']}")
            
            if retrospective.get('key_findings'):
                doc.add_paragraph('Key Findings from Historical Data:')
                for finding in retrospective['key_findings'][:5]:
                    doc.add_paragraph(f"{finding}", style='List Bullet')
            
            if retrospective.get('data_gaps'):
                doc.add_paragraph('Data Gaps Identified:')
                for gap in retrospective['data_gaps'][:5]:
                    p = doc.add_paragraph(f"⚠ {gap}", style='List Bullet')
        
        # Forward-Looking Analysis (Plan aspect)
        forward = hybrid_analysis.get('forward_looking_analysis', {})
        if forward:
            doc.add_heading('Forward-Looking Analysis (Future Plans)', 2)
            
            if forward.get('has_action_plan'):
                doc.add_paragraph("✓ Document includes an action plan")
            else:
                doc.add_paragraph("⚠ Document lacks a clear action plan")
            
            if forward.get('addresses_root_causes'):
                doc.add_paragraph("✓ Plans address root causes (hypoxia/DO)")
            else:
                doc.add_paragraph("⚠ Plans do not address root causes")
            
            if forward.get('timeline_present'):
                doc.add_paragraph("✓ Implementation timeline is present")
            else:
                doc.add_paragraph("⚠ No clear implementation timeline")
            
            if forward.get('proposed_interventions'):
                doc.add_paragraph('Proposed Interventions:')
                for intervention in forward['proposed_interventions'][:5]:
                    doc.add_paragraph(f"{intervention}", style='List Bullet')
            
            if forward.get('problematic_plans'):
                doc.add_paragraph('Concerning Proposals:')
                for issue in forward['problematic_plans'][:5]:
                    p = doc.add_paragraph(f"⚠ {issue}", style='List Bullet')
            
            if forward.get('recommended_improvements'):
                doc.add_paragraph('Recommended Improvements to Plans:')
                for improvement in forward['recommended_improvements'][:5]:
                    doc.add_paragraph(f"→ {improvement}", style='List Bullet')
        
        # AI Hybrid Insights (if available)
        if ai_insights and ai_insights.get('hybrid_analysis'):
            ai_hybrid = ai_insights['hybrid_analysis']
            if not ai_hybrid.get('error'):
                doc.add_heading('AI-Powered Alignment Assessment', 2)
                
                alignment = ai_hybrid.get('alignment_assessment', {})
                if alignment:
                    plans_match = alignment.get('plans_match_data', False)
                    align_score = alignment.get('alignment_score', 0)
                    
                    p = doc.add_paragraph()
                    p.add_run('Plans Match Data: ').bold = True
                    if plans_match:
                        match_run = p.add_run('YES')
                        match_run.font.color.rgb = self.brand_colors['success']
                    else:
                        match_run = p.add_run('NO - Plans may not address issues shown in data')
                        match_run.font.color.rgb = self.brand_colors['danger']
                    
                    doc.add_paragraph(f"Alignment Score: {align_score}/10")
                    
                    if alignment.get('gaps_between_data_and_plans'):
                        doc.add_paragraph('Gaps Between Data and Plans:')
                        for gap in alignment['gaps_between_data_and_plans'][:3]:
                            doc.add_paragraph(f"⚠ {gap}", style='List Bullet')
                
                if ai_hybrid.get('executive_summary'):
                    doc.add_paragraph()
                    summary_p = doc.add_paragraph()
                    summary_p.add_run('AI Summary: ').bold = True
                    summary_p.add_run(ai_hybrid['executive_summary'])
        
        doc.add_paragraph()
    
    def _add_score_card(self, doc: Document, compliance_eval: Dict):
        """Add detailed score card"""
        doc.add_heading('Compliance Score Card', 1)
        
        # Create table for scores
        table = doc.add_table(rows=5, cols=3)
        table.style = 'Light Grid'
        
        # Headers
        headers = table.rows[0].cells
        headers[0].text = 'Category'
        headers[1].text = 'Status'
        headers[2].text = 'Score Impact'
        
        # Make headers bold
        for cell in headers:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Critical Parameters
        row = table.rows[1].cells
        row[0].text = 'Critical Parameters'
        found = len(compliance_eval['critical_parameters']['found'])
        missing = len(compliance_eval['critical_parameters']['missing'])
        row[1].text = f"{found} found, {missing} missing"
        row[2].text = f"{found * 10 - missing * 10:+d} points"
        
        # Problematic Parameters
        row = table.rows[2].cells
        row[0].text = 'Problematic Parameters'
        prob_count = len(compliance_eval['problematic_parameters']['found'])
        row[1].text = f"{prob_count} found"
        row[2].text = f"{prob_count * -5:+d} points"
        
        # Calculations
        row = table.rows[3].cells
        row[0].text = 'Critical Calculations'
        calc_found = len(compliance_eval['calculations']['found'])
        calc_missing = len(compliance_eval['calculations']['missing'])
        row[1].text = f"{calc_found} performed, {calc_missing} missing"
        row[2].text = f"{calc_found * 15 - calc_missing * 15:+d} points"
        
        # Total
        row = table.rows[4].cells
        row[0].text = 'TOTAL SCORE'
        row[1].text = f"{compliance_eval['compliance_level'].value.upper()}"
        row[2].text = f"{compliance_eval['compliance_percentage']:.1f}%"
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        doc.add_paragraph()
    
    def _add_individual_parameter_comments(self, doc: Document, compliance_eval: Dict):
        """Add individual comments for each parameter as per Brief requirement"""
        doc.add_heading('Individual Parameter Analysis', 1)
        doc.add_paragraph(
            'Below is our assessment of each parameter in your Lake Management Plan:'
        )
        
        comments = compliance_eval.get('individual_parameter_comments', {})
        
        if comments:
            # Group by category
            critical_comments = {k: v for k, v in comments.items() if k.startswith('critical_')}
            problem_comments = {k: v for k, v in comments.items() if k.startswith('problem_')}
            calc_comments = {k: v for k, v in comments.items() if k.startswith('calc_')}
            
            # Critical parameters
            if critical_comments:
                doc.add_heading('Critical Parameters', 2)
                for param, comment in critical_comments.items():
                    p = doc.add_paragraph()
                    param_name = param.replace('critical_', '').replace('_', ' ').title()
                    p.add_run(f'• {param_name}: ').bold = True
                    p.add_run(comment)
            
            # Calculations
            if calc_comments:
                doc.add_heading('Critical Calculations', 2)
                for calc, comment in calc_comments.items():
                    p = doc.add_paragraph()
                    calc_name = calc.replace('calc_', '').replace('_', ' ').title()
                    p.add_run(f'• {calc_name}: ').bold = True
                    p.add_run(comment)
            
            # Problematic parameters
            if problem_comments:
                doc.add_heading('Non-Actionable Parameters', 2)
                for param, comment in problem_comments.items():
                    p = doc.add_paragraph()
                    param_name = param.replace('problem_', '').replace('_', ' ').title()
                    p.add_run(f'• {param_name}: ').bold = True
                    p.add_run(comment)
        
        doc.add_paragraph()
    
    def _add_lake_assessment(self, doc: Document, compliance_eval: Dict):
        """Add lake condition assessment as per Brief requirement"""
        doc.add_heading('Lake Condition Assessment', 1)
        
        assessment = compliance_eval.get('lake_assessment', {})
        
        if assessment:
            # Risk level with color coding
            risk = assessment.get('risk_level', 'moderate')
            p = doc.add_paragraph()
            p.add_run('Overall Risk Level: ').bold = True
            risk_run = p.add_run(risk.upper())
            risk_run.bold = True
            
            # Color based on risk level
            if risk in ['critical', 'high']:
                risk_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            elif risk == 'elevated':
                risk_run.font.color.rgb = RGBColor(255, 140, 0)  # Dark orange
            elif risk == 'moderate':
                risk_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
            elif risk == 'low':
                risk_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
            else:  # unknown or other
                risk_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
            
            # Hypoxia status
            hypoxia = assessment.get('hypoxia_status', 'not assessed')
            hab = assessment.get('hab_potential', 'not assessed')
            trajectory = assessment.get('trajectory', 'uncertain')
            data_qual = assessment.get('data_quality', 'not assessed')
            
            doc.add_paragraph(f"Hypoxia Status: {hypoxia.replace('_', ' ').title()}")
            doc.add_paragraph(f"HAB Potential: {hab.replace('_', ' ').title()}")
            doc.add_paragraph(f"Trajectory: {trajectory.replace('_', ' ').title()}")
            doc.add_paragraph(f"Data Quality: {data_qual.replace('_', ' ').title()}")
            
            # Key concerns
            if assessment.get('key_concerns'):
                doc.add_heading('Key Concerns', 2)
                for concern in assessment['key_concerns']:
                    doc.add_paragraph(f'{concern}', style='List Bullet')
            
            # Positive indicators
            if assessment.get('positive_indicators'):
                doc.add_heading('Positive Indicators', 2)
                for indicator in assessment['positive_indicators']:
                    doc.add_paragraph(f'{indicator}', style='List Bullet')
        else:
            doc.add_paragraph('Insufficient data to assess lake condition.')
        
        doc.add_paragraph()
    
    def _add_parameter_analysis(self, doc: Document, compliance_eval: Dict):
        """Add detailed parameter analysis"""
        doc.add_heading('Parameter Analysis', 1)
        
        # Critical Parameters
        doc.add_heading('Critical Parameters Assessment', 2)
        
        critical = compliance_eval['critical_parameters']
        
        if critical['found']:
            doc.add_paragraph('[YES] Parameters Measured:')
            for param in critical['found']:
                name = param['name'].replace('_', ' ').title()
                # Clean format per client feedback - just parameter and importance
                doc.add_paragraph(f"{name}", style='List Bullet')
                if param.get('importance'):
                    doc.add_paragraph(f"    -> {param['importance']}", style='Normal')
        
        if critical['missing']:
            doc.add_paragraph('[NO] Missing Critical Parameters:')
            for param in critical['missing']:
                name = param['name'].replace('_', ' ').title()
                doc.add_paragraph(f"{name}", style='List Bullet')
                doc.add_paragraph(f"    -> Why it matters: {param['importance']}", style='Normal')
                if param.get('requirements'):
                    doc.add_paragraph("    -> Requirements:", style='Normal')
                    for req in param['requirements'][:2]:
                        doc.add_paragraph(f"      - {req}", style='Normal')
        
        # Problematic Parameters
        if compliance_eval['problematic_parameters']['found']:
            doc.add_heading('Problematic Parameters Found', 2)
            doc.add_paragraph('The following non-actionable parameters were found:')
            
            for issue in compliance_eval['problematic_parameters']['issues']:
                param_name = issue['parameter'].replace('_', ' ').title()
                doc.add_paragraph(f"{param_name}: {issue['issue']}", style='List Bullet')
        
        doc.add_paragraph()
    
    def _add_calculations_analysis(self, doc: Document, compliance_eval: Dict):
        """Add calculations analysis section"""
        doc.add_heading('Critical Calculations', 1)
        
        calcs = compliance_eval['calculations']
        
        if calcs['found']:
            doc.add_paragraph('[YES] Calculations Performed:')
            for calc in calcs['found']:
                name = calc['name'].replace('_', ' ').title()
                doc.add_paragraph(f"{name}", style='List Bullet')
                if calc.get('importance'):
                    doc.add_paragraph(f"    -> {calc['importance']}", style='Normal')
        
        if calcs['missing']:
            doc.add_paragraph('[NO] Missing Calculations:')
            for calc in calcs['missing']:
                name = calc['name'].replace('_', ' ').title()
                doc.add_paragraph(f"{name}", style='List Bullet')
                if calc.get('formula'):
                    doc.add_paragraph(f"    -> Formula: {calc['formula']}", style='Normal')
                if calc.get('importance'):
                    doc.add_paragraph(f"    -> Why needed: {calc['importance']}", style='Normal')
        
        doc.add_paragraph()
    
    def _add_phytoplankton_management(self, doc: Document, compliance_eval: Dict):
        """Add phytoplankton management analysis section (per client feedback)"""
        phyto_mgmt = compliance_eval.get('phytoplankton_management', {})
        
        doc.add_heading('Phytoplankton Measurement Assessment', 1)
        
        # Phytoplankton measurement questions (per client feedback)
        doc.add_paragraph("We need to assess phytoplankton measurement:")
        
        questions = [
            "Are you actually identifying phytoplankton types and species?",
            "Are you calculating the percentage of:"
        ]
        
        for question in questions:
            doc.add_paragraph(f'{question}', style='List Bullet')
        
        # Sub-items for percentage calculation
        doc.add_paragraph('    ○ Cyanobacteria')
        doc.add_paragraph('        ■ Toxin producers')
        doc.add_paragraph('        ■ Non-Toxin producers')
        doc.add_paragraph('    ○ Beneficial algae')
        
        doc.add_paragraph('Are you measuring total phytoplankton biovolume?', style='List Bullet')
        
        doc.add_paragraph()
        
        # Chemical interventions assessment
        interventions_found = phyto_mgmt.get('interventions_found', [])
        is_negative = phyto_mgmt.get('is_negative', False)
        
        doc.add_heading('Phytoplankton Management Interventions', 2)
        
        if not interventions_found:
            # Check text for management approaches
            doc.add_paragraph(
                "Need to assess what is being done to manage phytoplankton:"
            )
            doc.add_paragraph("Algaecides", style='List Bullet')
            doc.add_paragraph("Herbicides", style='List Bullet')
            doc.add_paragraph("Phosphorus precipitants", style='List Bullet')
            doc.add_paragraph()
            
            p = doc.add_paragraph()
            p.add_run("Note: ").bold = True
            p.add_run("If these chemical interventions are being used, it is a significant negative. ")
            p.add_run("These approaches treat symptoms, not root causes.")
        else:
            # Interventions were detected
            p = doc.add_paragraph()
            p.add_run("[!] CONCERNING INTERVENTIONS DETECTED:").bold = True
            doc.add_paragraph()
            
            for intervention in interventions_found:
                name = intervention.get('name', 'Unknown')
                issue = intervention.get('issue', '')
                problems = intervention.get('problems', [])
                
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{name}: ").bold = True
                p.add_run(issue)
                
                for problem in problems:
                    doc.add_paragraph(f"    - {problem}", style='Normal')
            
            doc.add_paragraph()
            
            # Add warning
            warning_p = doc.add_paragraph()
            warning_p.add_run("Note: ").bold = True
            warning_p.add_run(
                "Chemical interventions like algaecides, herbicides, and phosphorus precipitants "
                "treat symptoms rather than root causes. Focus should be on reducing hypoxia."
            )
        
        doc.add_paragraph()
    
    def _add_ai_insights(self, doc: Document, ai_insights: Dict):
        """Add AI-generated insights section"""
        doc.add_heading('Advanced Analysis Insights', 1)
        
        # Focus Analysis
        if ai_insights.get('focus_analysis'):
            focus = ai_insights['focus_analysis']
            doc.add_heading('Focus Assessment', 2)
            
            p = doc.add_paragraph()
            p.add_run('Primary Focus: ').bold = True
            focus_type = focus.get('focus', 'unknown').replace('_', ' ').title()
            p.add_run(focus_type)
            
            if focus.get('root_cause_percentage'):
                doc.add_paragraph(
                    f"Root Cause Coverage: {focus['root_cause_percentage']}%"
                )
            
            if focus.get('key_findings'):
                doc.add_paragraph('Key Observations:')
                for finding in focus['key_findings'][:3]:
                    doc.add_paragraph(f"{finding}", style='List Bullet')
        
        # Overall Quality
        if ai_insights.get('overall_quality'):
            quality = ai_insights['overall_quality']
            doc.add_heading('Report Quality Assessment', 2)
            
            if quality.get('executive_summary'):
                doc.add_paragraph(quality['executive_summary'])
            
            if quality.get('major_gaps'):
                doc.add_paragraph('Major Gaps Identified:')
                for gap in quality['major_gaps'][:3]:
                    doc.add_paragraph(f"{gap}", style='List Bullet')
        
        doc.add_paragraph()
    
    def _add_recommendations(self, doc: Document, compliance_eval: Dict):
        """Add recommendations section with YouTube links as per Brief"""
        doc.add_heading('Recommendations to Convert to ACTION Plan', 1)
        
        recommendations = compliance_eval.get('recommendations', [])
        
        if not recommendations:
            doc.add_paragraph('No specific recommendations at this time.')
            return
        
        # Load educational resources from config
        from config import COMPLIANCE_RULES
        educational_links = COMPLIANCE_RULES.get('educational_resources', {})
        
        # Group by priority
        high_priority = [r for r in recommendations if r.get('priority') == 'HIGH']
        medium_priority = [r for r in recommendations if r.get('priority') == 'MEDIUM']
        low_priority = [r for r in recommendations if r.get('priority') == 'LOW']
        
        if high_priority:
            doc.add_heading('High Priority Actions', 2)
            for i, rec in enumerate(high_priority[:5], 1):
                p = doc.add_paragraph()
                p.add_run(f"{i}. {rec['recommendation']}").bold = True
                doc.add_paragraph(f"   {rec.get('explanation', '')}")
                if rec.get('requirements'):
                    for req in rec['requirements'][:2]:
                        doc.add_paragraph(f"   - {req}")
                
                # Add YouTube link if available
                param_key = rec.get('recommendation', '').lower()
                for edu_key, url in educational_links.items():
                    if edu_key in param_key or edu_key.replace('_', ' ') in param_key.lower():
                        doc.add_paragraph(f"   [VIDEO] Learn more: {url}")
                        break
        
        if medium_priority:
            doc.add_heading('Medium Priority Improvements', 2)
            for rec in medium_priority[:3]:
                doc.add_paragraph(rec['recommendation'], style='List Bullet')
                if rec.get('explanation'):
                    doc.add_paragraph(f"  {rec['explanation']}")
        
        if low_priority:
            doc.add_heading('Additional Considerations', 2)
            for rec in low_priority[:3]:
                doc.add_paragraph(rec['recommendation'], style='List Bullet')
        
        # Add monthly tracking recommendation from Brief
        doc.add_paragraph()
        doc.add_paragraph(
            "Note: All parameters should ideally be measured monthly and tracked over time "
            "to identify trends and seasonal patterns."
        )
        
        # Add YouTube Educational Resources section
        doc.add_heading('Educational Resources', 2)
        doc.add_paragraph(
            "Learn more about why these recommendations matter:"
        )
        from config import COMPLIANCE_RULES
        edu_links = COMPLIANCE_RULES.get('educational_resources', {})
        
        doc.add_paragraph("YouTube Videos:")
        doc.add_paragraph("Dissolved Oxygen Dynamics: " + edu_links.get('dissolved_oxygen', ''), style='List Bullet')
        doc.add_paragraph("Phytoplankton Analysis: " + edu_links.get('phytoplankton_analysis', ''), style='List Bullet')
        doc.add_paragraph("Bathymetry Calculations: " + edu_links.get('bathymetry_calculations', ''), style='List Bullet')
        doc.add_paragraph("Monthly Tracking Importance: " + edu_links.get('monthly_tracking', ''), style='List Bullet')
        doc.add_paragraph("Hypoxia and HABs: " + edu_links.get('hypoxia_hab_link', ''), style='List Bullet')
        
        doc.add_paragraph()
    
    def _add_call_to_action(self, doc: Document):
        """Add call to action section with constructive language"""
        doc.add_page_break()
        doc.add_heading('CALL TO ACTION: Transform Your Lake Management Plan into a Lake Management ACTION Plan', 1)
        
        # Softened opening statement (per client feedback)
        p = doc.add_paragraph()
        p.add_run(
            "Your current Lake Management Plan is measuring many non-actionable parameters. "
        ).bold = True
        p.add_run(
            "Several other key parameters are not being measured. "
            "By reallocating resources to the right measurements, you can achieve better outcomes."
        )
        
        doc.add_paragraph()
        
        doc.add_heading('Why This Matters', 2)
        doc.add_paragraph(
            "The GAO Report (GAO-22-104449) confirms what our analysis shows: "
            "hypoxia and HABs are inextricably linked. By focusing on root causes, you can:"
        )
        doc.add_paragraph('Better understand your lake\'s true condition', style='List Bullet')
        doc.add_paragraph('Make more informed management decisions', style='List Bullet')
        doc.add_paragraph('Achieve measurable improvements in water quality', style='List Bullet')
        doc.add_paragraph('Protect property values around the lake', style='List Bullet')
        doc.add_paragraph('Ensure safe recreational use', style='List Bullet')
        
        doc.add_heading('Your Path to Lake Recovery', 2)
        doc.add_paragraph('Reallocate budget from non-actionable parameters to critical measurements', style='List Number')
        doc.add_paragraph('Begin monthly monitoring of DO profiles and hypoxic volumes', style='List Number')
        doc.add_paragraph('Implement bathymetry-based calculations', style='List Number')
        doc.add_paragraph('Focus interventions on reducing hypoxia', style='List Number')
        doc.add_paragraph('Track phytoplankton taxonomy and biovolume monthly', style='List Number')
        
        doc.add_heading('Convert to a Lake Management ACTION Plan Today', 2)
        action_p = doc.add_paragraph()
        action_p.add_run(
            "Take the next step toward better lake management. "
        ).bold = True
        action_p.add_run(
            "The science is clear (GAO-22-104449): hypoxia drives HABs. Address hypoxia, improve your lake. "
            "This analysis compared your report against OUR THINKING - proven best practices that work."
        )
        
        doc.add_paragraph()
        
        # CTA box
        cta = doc.add_paragraph()
        
        cta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cta.add_run('Schedule Your ACTION Plan Consultation\n').bold = True
        cta.add_run('Call: 1-800-LAKE-FIX\n').bold = True
        cta.add_run('Email: action@reporttoreveal.com\n').bold = True
        cta.add_run('Visit: www.reporttoreveal.com/action\n').bold = True
        
        doc.add_paragraph()
        
        # Encouraging closing
        final = doc.add_paragraph()
        final.alignment = WD_ALIGN_PARAGRAPH.CENTER
        final.add_run('Every improvement starts with better data.\n').italic = True
        final.add_run('Let us help you make the change.').bold = True
    
    def _add_appendix(self, doc: Document, doc_data: Dict, compliance_eval: Dict):
        """Add technical appendix"""
        doc.add_page_break()
        doc.add_heading('Appendix: Technical Details', 1)
        
        doc.add_heading('Parameters Detected', 2)
        
        # Create a simple list of all parameters found
        params = doc_data.get('parameters_found', {})
        
        if params:
            for param, found in params.items():
                status = "[YES]" if found else "[NO]"
                clean_name = param.replace('critical_', '').replace('problem_', '').replace('calc_', '')
                clean_name = clean_name.replace('_', ' ').title()
                doc.add_paragraph(f"{status} {clean_name}")
        
        # Add metrics if found
        if doc_data.get('metrics'):
            doc.add_heading('Extracted Metrics', 2)
            metrics = doc_data['metrics']
            
            for metric_type, values in metrics.items():
                if values:
                    clean_type = metric_type.replace('_', ' ').title()
                    doc.add_paragraph(f"{clean_type}: {values}")
        
        # Add analysis metadata
        doc.add_heading('Analysis Metadata', 2)
        doc.add_paragraph(f"Document Pages: {doc_data.get('page_count', 'N/A')}")
        doc.add_paragraph(f"Text Extraction Method: {doc_data.get('extraction_method', 'N/A')}")
        doc.add_paragraph(f"Has Tables: {doc_data.get('has_tables', False)}")
        doc.add_paragraph(f"Has Images: {doc_data.get('has_images', False)}")
        doc.add_paragraph(f"Analysis Version: 1.0")
        doc.add_paragraph(f"Compliance Rules Version: {datetime.now().strftime('%Y-%m')}")
